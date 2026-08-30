"""Turn a model answer (Markdown) into a real document.

An answer worth keeping is usually an answer somebody has to send on — to a
supervisor, into a report, to a client. Copying markdown out of a chat window
and re-formatting it by hand is the last manual step in a tool whose whole point
is removing manual steps, so OMNIX renders the document itself.

Five formats, each with a reason to exist:

    docx   the one people actually need — editable, tracked-changes-able
    pdf    the one people send when it must not be edited
    html   self-contained, for pasting into a wiki or an email
    md     the source, for anyone who wants it back in a tool
    txt    plain, for when formatting is the problem

The markdown subset handled here is exactly what the models emit and what
`components/Markdown.tsx` renders: headings, bold/italic/inline code, fenced
code, bullet and ordered lists, pipe tables, block quotes, rules and links. That
symmetry is deliberate — a document that silently drops a table the screen
showed would be worse than no export at all.

Both writers consume ONE parse (`parse_blocks`), so DOCX and PDF cannot drift
apart in what they understand.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from datetime import datetime

FORMATS = ("docx", "pdf", "html", "md", "txt")

MIME = {
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
    "html": "text/html; charset=utf-8",
    "md": "text/markdown; charset=utf-8",
    "txt": "text/plain; charset=utf-8",
}


# ---------------------------------------------------------------------------
# Parsing
# ---------------------------------------------------------------------------
@dataclass
class Block:
    kind: str                       # h1..h6 | p | code | ul | ol | table | quote | rule
    text: str = ""
    lang: str = ""
    items: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)


_TABLE_ROW = re.compile(r"^\s*\|.*\|\s*$")
_DIVIDER = re.compile(r"^\s*\|?[\s:\-]*-[\s:|\-]*\|?\s*$")


def _cells(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def parse_blocks(md: str) -> list[Block]:
    """Markdown -> a flat block list. Deliberately not a full CommonMark parser;
    it handles what the models emit and passes anything else through as text."""
    lines = (md or "").replace("\r\n", "\n").split("\n")
    out: list[Block] = []
    para: list[str] = []
    i = 0

    def flush() -> None:
        nonlocal para
        if para:
            out.append(Block("p", " ".join(para)))
            para = []

    while i < len(lines):
        line = lines[i]
        t = line.strip()

        if t.startswith("```"):
            flush()
            lang = t[3:].strip().split()[0] if len(t) > 3 else ""
            buf: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1                                  # step over the closing fence
            out.append(Block("code", "\n".join(buf), lang=lang))
            continue

        if not t:
            flush()
            i += 1
            continue

        if re.fullmatch(r"([-*_])\1{2,}", t):
            flush()
            out.append(Block("rule"))
            i += 1
            continue

        head = re.match(r"^(#{1,6})\s+(.*)$", t)
        if head:
            flush()
            out.append(Block(f"h{len(head.group(1))}", head.group(2).strip()))
            i += 1
            continue

        if _TABLE_ROW.match(line) and i + 1 < len(lines) and _DIVIDER.match(lines[i + 1]):
            flush()
            rows = [_cells(line)]
            i += 2
            while i < len(lines) and _TABLE_ROW.match(lines[i]):
                rows.append(_cells(lines[i]))
                i += 1
            out.append(Block("table", rows=rows))
            continue

        if t.startswith(">"):
            flush()
            buf = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                buf.append(re.sub(r"^>\s?", "", lines[i].strip()))
                i += 1
            out.append(Block("quote", " ".join(buf)))
            continue

        bullet = re.match(r"^[-*+]\s+(.*)$", t)
        ordered = re.match(r"^\d+[.)]\s+(.*)$", t)
        if bullet or ordered:
            flush()
            is_ord = bool(ordered)
            items: list[str] = []
            while i < len(lines):
                s = lines[i].strip()
                b = re.match(r"^[-*+]\s+(.*)$", s)
                o = re.match(r"^\d+[.)]\s+(.*)$", s)
                if (o if is_ord else b):
                    items.append((o if is_ord else b).group(1))
                    i += 1
                    continue
                # A wrapped continuation line belongs to the item above it.
                if s and not re.match(r"^([-*+]|\d+[.)])\s", s) and items:
                    items[-1] += " " + s
                    i += 1
                    continue
                break
            out.append(Block("ol" if is_ord else "ul", items=items))
            continue

        para.append(t)
        i += 1

    flush()
    return out


# Inline spans, longest delimiter first so `**` is not eaten by `*`.
_INLINE = re.compile(
    r"(`[^`]+`)|(\*\*[^*]+\*\*)|(__[^_]+__)|(\*[^*]+\*)|(_[^_]+_)"
    r"|(\[[^\]]+\]\([^)]+\))")


def inline_runs(text: str) -> list[tuple[str, str]]:
    """Split a line into (text, style) where style is '', 'b', 'i' or 'code'.
    Links become their label plus the URL in parentheses — a printed document
    cannot be clicked, so a bare label would lose the reference entirely."""
    runs: list[tuple[str, str]] = []
    last = 0
    for m in _INLINE.finditer(text):
        if m.start() > last:
            runs.append((text[last:m.start()], ""))
        tok = m.group(0)
        if tok.startswith("`"):
            runs.append((tok[1:-1], "code"))
        elif tok.startswith("**") or tok.startswith("__"):
            runs.append((tok[2:-2], "b"))
        elif tok.startswith("["):
            cut = tok.index("](")
            runs.append((tok[1:cut], "b"))
            runs.append((f" ({tok[cut + 2:-1]})", "i"))
        else:
            runs.append((tok[1:-1], "i"))
        last = m.end()
    if last < len(text):
        runs.append((text[last:], ""))
    return runs or [(text, "")]


def strip_markdown(md: str) -> str:
    """Plain text, for the .txt export and as a TTS fallback."""
    out: list[str] = []
    for b in parse_blocks(md):
        if b.kind.startswith("h"):
            out.append(b.text.upper())
        elif b.kind == "code":
            out.append(b.text)
        elif b.kind in ("ul", "ol"):
            for n, it in enumerate(b.items, 1):
                prefix = f"{n}. " if b.kind == "ol" else "- "
                out.append(prefix + "".join(t for t, _ in inline_runs(it)))
        elif b.kind == "table":
            for row in b.rows:
                out.append("\t".join(row))
        elif b.kind == "rule":
            out.append("-" * 40)
        else:
            out.append("".join(t for t, _ in inline_runs(b.text)))
        out.append("")
    return "\n".join(out).strip()


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def to_docx(md: str, title: str, subtitle: str = "") -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(7)

    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.9)
        s.left_margin = s.right_margin = Inches(1.0)

    p = doc.add_paragraph()
    r = p.add_run(title)
    r.font.size = Pt(22)
    r.font.bold = True
    p.paragraph_format.space_after = Pt(2)

    sub = subtitle or datetime.now().strftime("Generated by OMNIX · %d %B %Y, %H:%M")
    p = doc.add_paragraph()
    r = p.add_run(sub)
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    p.paragraph_format.space_after = Pt(16)

    def runs_into(par, text: str) -> None:
        for chunk, style in inline_runs(text):
            if not chunk:
                continue
            run = par.add_run(chunk)
            run.font.bold = style == "b"
            run.font.italic = style == "i"
            if style == "code":
                run.font.name = "Consolas"
                run.font.size = Pt(10)

    for b in parse_blocks(md):
        if b.kind.startswith("h"):
            level = min(4, int(b.kind[1]))
            par = doc.add_paragraph()
            par.paragraph_format.space_before = Pt(12)
            par.paragraph_format.space_after = Pt(4)
            run = par.add_run(b.text)
            run.font.bold = True
            run.font.size = Pt({1: 17, 2: 14, 3: 12.5, 4: 11.5}[level])
        elif b.kind == "code":
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Inches(0.25)
            par.paragraph_format.space_after = Pt(10)
            run = par.add_run(b.text)
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        elif b.kind in ("ul", "ol"):
            for it in b.items:
                par = doc.add_paragraph(
                    style="List Number" if b.kind == "ol" else "List Bullet")
                par.paragraph_format.space_after = Pt(3)
                runs_into(par, it)
        elif b.kind == "table" and b.rows:
            cols = max(len(r_) for r_ in b.rows)
            t = doc.add_table(rows=0, cols=cols)
            t.style = "Light Grid Accent 1"
            for ri, row in enumerate(b.rows):
                cells = t.add_row().cells
                for ci in range(cols):
                    cells[ci].text = ""
                    par = cells[ci].paragraphs[0]
                    val = row[ci] if ci < len(row) else ""
                    if ri == 0:
                        run = par.add_run(val)
                        run.font.bold = True
                        run.font.size = Pt(10)
                    else:
                        runs_into(par, val)
            doc.add_paragraph()
        elif b.kind == "quote":
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Inches(0.35)
            runs_into(par, b.text)
            for run in par.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif b.kind == "rule":
            par = doc.add_paragraph("─" * 40)
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            runs_into(doc.add_paragraph(), b.text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def to_pdf(md: str, title: str, subtitle: str = "") -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (KeepTogether, Paragraph, SimpleDocTemplate,
                                    Spacer, Table, TableStyle)

    ss = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=ss["BodyText"], fontName="Helvetica",
                          fontSize=10, leading=15, spaceAfter=7, alignment=TA_LEFT)
    mono = ParagraphStyle("mono", parent=body, fontName="Courier", fontSize=8.5,
                          leading=12, leftIndent=10, backColor=colors.HexColor("#F4F4F5"),
                          borderPadding=6, spaceAfter=10)
    quote = ParagraphStyle("quote", parent=body, leftIndent=14,
                           textColor=colors.HexColor("#555555"),
                           fontName="Helvetica-Oblique")
    heads = {
        1: ParagraphStyle("h1", parent=body, fontSize=16, leading=20,
                          spaceBefore=12, spaceAfter=5, fontName="Helvetica-Bold"),
        2: ParagraphStyle("h2", parent=body, fontSize=13, leading=17,
                          spaceBefore=11, spaceAfter=4, fontName="Helvetica-Bold"),
        3: ParagraphStyle("h3", parent=body, fontSize=11.5, leading=15,
                          spaceBefore=9, spaceAfter=3, fontName="Helvetica-Bold"),
        4: ParagraphStyle("h4", parent=body, fontSize=10.5, leading=14,
                          spaceBefore=8, spaceAfter=3, fontName="Helvetica-Bold"),
    }

    def esc(s: str) -> str:
        return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

    def rich(text: str) -> str:
        """Inline runs as reportlab's mini-HTML. Everything is escaped FIRST, so
        a model writing `<script>` or a stray `&` cannot break the parser."""
        parts = []
        for chunk, style in inline_runs(text):
            e = esc(chunk)
            if style == "b":
                parts.append(f"<b>{e}</b>")
            elif style == "i":
                parts.append(f"<i>{e}</i>")
            elif style == "code":
                parts.append(f'<font face="Courier" size="9">{e}</font>')
            else:
                parts.append(e)
        return "".join(parts)

    story = [
        Paragraph(esc(title), ParagraphStyle(
            "title", parent=body, fontSize=21, leading=25,
            fontName="Helvetica-Bold", spaceAfter=2)),
        Paragraph(esc(subtitle or datetime.now().strftime(
            "Generated by OMNIX · %d %B %Y, %H:%M")),
            ParagraphStyle("sub", parent=body, fontSize=9,
                           textColor=colors.HexColor("#777777"), spaceAfter=16)),
    ]

    for b in parse_blocks(md):
        if b.kind.startswith("h"):
            story.append(Paragraph(rich(b.text), heads[min(4, int(b.kind[1]))]))
        elif b.kind == "code":
            # <br/> rather than real newlines: reportlab collapses whitespace,
            # so a code block posted verbatim comes out as one long line.
            text = esc(b.text).replace(" ", "&nbsp;").replace("\n", "<br/>")
            story.append(Paragraph(text, mono))
        elif b.kind in ("ul", "ol"):
            for n, it in enumerate(b.items, 1):
                marker = f"{n}." if b.kind == "ol" else "•"
                story.append(Paragraph(
                    f"{marker}&nbsp;&nbsp;{rich(it)}",
                    ParagraphStyle("li", parent=body, leftIndent=14,
                                   spaceAfter=3)))
        elif b.kind == "table" and b.rows:
            cols = max(len(r) for r in b.rows)
            cell = ParagraphStyle("cell", parent=body, fontSize=8.5, leading=11,
                                  spaceAfter=0)
            head = ParagraphStyle("cellh", parent=cell,
                                  fontName="Helvetica-Bold")
            data = [[Paragraph(rich(r[c] if c < len(r) else ""),
                               head if ri == 0 else cell)
                     for c in range(cols)]
                    for ri, r in enumerate(b.rows)]
            avail = A4[0] - 40 * mm
            t = Table(data, colWidths=[avail / cols] * cols, repeatRows=1)
            t.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EFEFEF")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(KeepTogether(t))
            story.append(Spacer(1, 9))
        elif b.kind == "quote":
            story.append(Paragraph(rich(b.text), quote))
        elif b.kind == "rule":
            story.append(Spacer(1, 5))
        else:
            story.append(Paragraph(rich(b.text), body))

    buf = io.BytesIO()
    SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm,
        topMargin=18 * mm, bottomMargin=18 * mm,
        title=title, author="OMNIX",
    ).build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# HTML
# ---------------------------------------------------------------------------
def to_html(md: str, title: str, subtitle: str = "") -> bytes:
    def esc(s: str) -> str:
        return (s.replace("&", "&amp;").replace("<", "&lt;")
                 .replace(">", "&gt;").replace('"', "&quot;"))

    def rich(text: str) -> str:
        parts = []
        for chunk, style in inline_runs(text):
            e = esc(chunk)
            parts.append({"b": f"<strong>{e}</strong>", "i": f"<em>{e}</em>",
                          "code": f"<code>{e}</code>"}.get(style, e))
        return "".join(parts)

    out: list[str] = []
    for b in parse_blocks(md):
        if b.kind.startswith("h"):
            n = min(6, int(b.kind[1]))
            out.append(f"<h{n}>{rich(b.text)}</h{n}>")
        elif b.kind == "code":
            cls = f' class="language-{esc(b.lang)}"' if b.lang else ""
            out.append(f"<pre><code{cls}>{esc(b.text)}</code></pre>")
        elif b.kind in ("ul", "ol"):
            tag = b.kind
            items = "".join(f"<li>{rich(i)}</li>" for i in b.items)
            out.append(f"<{tag}>{items}</{tag}>")
        elif b.kind == "table" and b.rows:
            head = "".join(f"<th>{rich(c)}</th>" for c in b.rows[0])
            rows = "".join(
                "<tr>" + "".join(f"<td>{rich(c)}</td>" for c in r) + "</tr>"
                for r in b.rows[1:])
            out.append(f"<table><thead><tr>{head}</tr></thead>"
                       f"<tbody>{rows}</tbody></table>")
        elif b.kind == "quote":
            out.append(f"<blockquote>{rich(b.text)}</blockquote>")
        elif b.kind == "rule":
            out.append("<hr>")
        else:
            out.append(f"<p>{rich(b.text)}</p>")

    sub = subtitle or datetime.now().strftime(
        "Generated by OMNIX · %d %B %Y, %H:%M")
    # Self-contained and print-ready: no external stylesheet, no font fetch, and
    # a light ground regardless of theme because this file gets emailed and
    # printed, not viewed inside the app.
    doc = f"""<!doctype html>
<html lang="en"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{esc(title)}</title>
<style>
  :root {{ color-scheme: light; }}
  body {{ max-width: 46rem; margin: 3rem auto; padding: 0 1.25rem;
    font: 16px/1.65 -apple-system, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
    color: #1b1b1f; background: #fff; }}
  h1.doc {{ font-size: 2rem; margin: 0 0 .2rem; letter-spacing: -.02em; }}
  .sub {{ color: #6b6b73; font-size: .85rem; margin: 0 0 2rem; }}
  h1,h2,h3,h4 {{ line-height: 1.25; margin: 1.8em 0 .5em; }}
  h2 {{ font-size: 1.35rem; }} h3 {{ font-size: 1.12rem; }}
  p, li {{ margin: 0 0 .7em; }}
  code {{ font: .88em/1.5 ui-monospace, Consolas, monospace;
    background: #f2f2f4; padding: .12em .35em; border-radius: 4px; }}
  pre {{ background: #f7f7f9; border: 1px solid #e6e6ea; border-radius: 8px;
    padding: .9rem 1rem; overflow-x: auto; }}
  pre code {{ background: none; padding: 0; font-size: .84rem; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1rem 0;
    font-size: .9rem; }}
  th, td {{ border: 1px solid #dcdce2; padding: .45rem .6rem;
    text-align: left; vertical-align: top; }}
  th {{ background: #f2f2f4; }}
  blockquote {{ margin: 1rem 0; padding: .1rem 0 .1rem 1rem;
    border-left: 3px solid #dcdce2; color: #55555d; }}
  hr {{ border: 0; border-top: 1px solid #e6e6ea; margin: 2rem 0; }}
  @media print {{ body {{ margin: 0; max-width: none; }} }}
</style></head><body>
<h1 class="doc">{esc(title)}</h1><p class="sub">{esc(sub)}</p>
{chr(10).join(out)}
</body></html>"""
    return doc.encode("utf-8")


# ---------------------------------------------------------------------------
def render(md: str, fmt: str, title: str, subtitle: str = "") -> bytes:
    fmt = (fmt or "docx").lower().strip()
    if fmt == "docx":
        return to_docx(md, title, subtitle)
    if fmt == "pdf":
        return to_pdf(md, title, subtitle)
    if fmt == "html":
        return to_html(md, title, subtitle)
    if fmt == "txt":
        # Titled like every other format. Without this a whitespace-only answer
        # produced a zero-byte file, which reads as a broken download rather
        # than as an empty document.
        head = f"{title}\n{'=' * len(title)}\n"
        if subtitle:
            head += f"{subtitle}\n"
        return f"{head}\n{strip_markdown(md)}".rstrip().encode("utf-8") + b"\n"
    if fmt == "md":
        return (f"# {title}\n\n_{subtitle}_\n\n{md}" if subtitle
                else f"# {title}\n\n{md}").encode("utf-8")
    raise ValueError(f"unsupported format: {fmt}")


_SAFE = re.compile(r"[^\w\s-]")


def filename(title: str, fmt: str) -> str:
    stem = _SAFE.sub(" ", title or "omnix").strip().lower()
    # Collapse runs: stripping punctuation out of 'Report: "Q3" / final' leaves
    # gaps that would otherwise become 'report-q3--final'.
    stem = re.sub(r"[\s-]+", "-", stem).strip("-")
    return f"{(stem or 'omnix')[:60].rstrip('-')}.{fmt}"
